from docx import Document
from textwrap import wrap
from sentence_transformers import SentenceTransformer
import chromadb
import requests
import csv
import os
import re

# Step 1: Read the Word document
doc = Document("Customer_Service_QA_Knowledge_Base.docx")

text = ""

for para in doc.paragraphs:
    if para.text.strip():
        text += para.text.strip() + "\n"

# This reads the Q&A tables inside the Word document
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            text += " | ".join(row_text) + "\n"

# Step 2: Split the text into chunks
chunks = wrap(text, 500)

print("Number of chunks created:", len(chunks))

# Step 3: Convert chunks into embeddings
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
embeddings = embedding_model.encode(chunks).tolist()

# Step 4: Store embeddings in ChromaDB
client = chromadb.Client()

try:
    client.delete_collection(name="customer_service_qa")
except:
    pass

collection = client.create_collection(name="customer_service_qa")

collection.add(
    documents=chunks,
    embeddings=embeddings,
    ids=[f"chunk_{i}" for i in range(len(chunks))]
)

print("ChromaDB vector database created successfully.")
print("Number of chunks stored:", collection.count())


# Step 5: Save lead/query details into CSV
def save_to_csv(user_query, lead_type, extracted_details, bot_response):
    file_name = "lead_queries.csv"
    file_exists = os.path.isfile(file_name)

    with open(file_name, mode="a", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)

        if not file_exists:
            writer.writerow([
                "User Query",
                "Lead Type",
                "Extracted Details",
                "Bot Response"
            ])

        writer.writerow([
            user_query,
            lead_type,
            extracted_details,
            bot_response
        ])


# Step 6: Classify the user query
def classify_lead(user_query):
    query = user_query.lower()

    support_keywords = [
        "password", "refund", "return", "login", "error", "issue",
        "problem", "track", "delivery", "cancel", "damaged",
        "order", "payment declined", "charged twice", "app crashing"
    ]

    hot_keywords = [
        "buy", "purchase", "pricing", "quote", "demo", "contact me",
        "sales", "subscription", "interested in buying", "need this",
        "book a call", "speak to sales"
    ]

    warm_keywords = [
        "interested", "looking for", "considering", "features",
        "compare", "more information", "tell me more", "options"
    ]

    if any(word in query for word in support_keywords):
        return "Support Query"

    if any(word in query for word in hot_keywords):
        return "Hot Lead"

    if any(word in query for word in warm_keywords):
        return "Warm Lead"

    return "Cold Lead"


# Step 7: Extract simple lead details from the user query
def extract_details(user_query):
    query = user_query.lower()

    details = {
        "Product Interest": "Not mentioned",
        "Team Size": "Not mentioned",
        "Budget": "Not mentioned",
        "Timeline": "Not mentioned",
        "Contact Request": "No"
    }

    # Product interest
    product_keywords = ["subscription", "software", "app", "product", "service", "plan"]
    for word in product_keywords:
        if word in query:
            details["Product Interest"] = word
            break

    # Team size / number of users
    team_match = re.search(r"(\d+)\s*(users|employees|people|staff|members|licenses)", query)
    if team_match:
        details["Team Size"] = team_match.group(0)

    # Budget
    budget_match = re.search(r"(\$|aed|dhs|£|gbp)?\s?\d+[,\d]*(\s?(per month|monthly|budget)?)", query)
    if budget_match:
        details["Budget"] = budget_match.group(0)

    # Timeline
    timelines = ["today", "this week", "next week", "this month", "next month", "as soon as possible", "asap"]
    for time in timelines:
        if time in query:
            details["Timeline"] = time
            break

    # Contact request
    contact_keywords = ["contact me", "call me", "email me", "speak to sales", "book a call", "schedule a demo"]
    if any(word in query for word in contact_keywords):
        details["Contact Request"] = "Yes"

    return details


# Step 8: Ask the chatbot using RAG + Ollama
def ask_chatbot(question):
    question_embedding = embedding_model.encode([question]).tolist()

    results = collection.query(
        query_embeddings=question_embedding,
        n_results=3
    )

    retrieved_chunks = results["documents"][0]
    context = "\n\n".join(retrieved_chunks)

    prompt = f"""
You are a helpful customer service assistant.

Use only the information from the context below to answer the user's question.
If the answer is not available in the context, say that you do not have enough information.

Context:
{context}

User Question:
{question}

Answer:
"""

    response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": "llama3.2:1b",
            "prompt": prompt,
            "stream": False
        }
    )

    return response.json()["response"]


# Step 9: Simple terminal chat
print("\nCustomer Service RAG Chatbot with Lead Qualification is ready.")
print("Type 'exit' to stop.\n")

while True:
    user_question = input("Ask a question: ")

    if user_question.lower() == "exit":
        print("Chatbot stopped.")
        break

    answer = ask_chatbot(user_question)
    lead_type = classify_lead(user_question)
    extracted_details = extract_details(user_question)

    save_to_csv(
        user_question,
        lead_type,
        extracted_details,
        answer
    )

    print("\nLead Type:")
    print(lead_type)

    print("\nExtracted Details:")
    print(extracted_details)

    print("\nAnswer:")
    print(answer)

    print("\nSaved to lead_queries.csv")
    print("-" * 80)
