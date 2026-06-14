from docx import Document
from sentence_transformers import SentenceTransformer
import chromadb
import os
import shutil

DOC_PATH = "Customer_Service_QA_Knowledge_Base.docx"
CHROMA_PATH = "./chroma_db"
COLLECTION_NAME = "customer_service_qa"


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    chunks = []

    # Read normal paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)

    # Read tables and keep each Q&A row together
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            if len(cells) >= 2:
                question = cells[0]
                answer = cells[1]

                if question.lower() != "customer question":
                    chunks.append(f"Customer Question: {question}\nAgent Response: {answer}")

    return chunks


def create_semantic_chunks(raw_chunks, max_length=700):
    final_chunks = []
    current_chunk = ""

    for chunk in raw_chunks:
        if len(current_chunk) + len(chunk) <= max_length:
            current_chunk += chunk + "\n\n"
        else:
            if current_chunk.strip():
                final_chunks.append(current_chunk.strip())
            current_chunk = chunk + "\n\n"

    if current_chunk.strip():
        final_chunks.append(current_chunk.strip())

    return final_chunks


def create_vectordb():
    print("Reading document...")
    raw_chunks = extract_text_from_docx(DOC_PATH)

    print("Creating better chunks...")
    chunks = create_semantic_chunks(raw_chunks)

    print("Number of chunks created:", len(chunks))

    print("Loading embedding model...")
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

    print("Creating embeddings...")
    embeddings = embedding_model.encode(chunks).tolist()

    # Recreate persistent vector database
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)

    client = chromadb.PersistentClient(path=CHROMA_PATH)

    collection = client.create_collection(name=COLLECTION_NAME)

    collection.add(
        documents=chunks,
        embeddings=embeddings,
        ids=[f"chunk_{i}" for i in range(len(chunks))]
    )

    print("Persistent ChromaDB vector database created successfully.")
    print("Number of chunks stored:", collection.count())
    print("VectorDB saved in folder:", CHROMA_PATH)


if __name__ == "__main__":
    create_vectordb()